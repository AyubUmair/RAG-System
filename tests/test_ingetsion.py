import fitz
from pathlib import Path
from src.ingestion.parser import DocumentIngester
from src.config import settings
import pytest

def test_parse_text_and_chunking():
    """Verify raw text parsing and recursive chunk splitting."""
    ingester = DocumentIngester(chunk_size=100, chunk_overlap=20)
    
    sample_text = (
        "Agentic RAG combines retrieval mechanisms with autonomous decision-making agents. "
        "Unlike naive RAG systems, it evaluates retrieved context relevance dynamically. "
        "If the context is irrelevant, it reformulates the query to search again."
    )
    
    docs = ingester.parse_text(sample_text, source_name="agentic_overview.txt")
    assert len(docs) == 1
    assert docs[0]["metadata"]["source"] == "agentic_overview.txt"
    
    chunks = ingester.chunk_documents(docs)
    assert len(chunks) > 1
    
    # Check metadata propagation
    for idx, chunk in enumerate(chunks):
        assert "source" in chunk["metadata"]
        assert "page" in chunk["metadata"]
        assert chunk["metadata"]["chunk_id"] == f"agentic_overview.txt_p1_c{idx}"
        assert len(chunk["text"]) <= 120  # accounts for boundary words

def test_parse_pdf(tmp_path: Path):
    """Create a temporary PDF and verify page-by-page extraction."""
    # Generate a lightweight test PDF using PyMuPDF
    pdf_path = tmp_path / "sample_test.pdf"
    doc = fitz.open()
    
    # Page 1
    page1 = doc.new_page()
    page1.insert_text((50, 50), "Page 1: Introduction to Hybrid Vector Search.")
    
    # Page 2
    page2 = doc.new_page()
    page2.insert_text((50, 50), "Page 2: Cross-encoder reranking improves retrieval precision.")
    
    doc.save(str(pdf_path))
    doc.close()
    
    # Ingest the generated PDF
    ingester = DocumentIngester(chunk_size=settings.CHUNK_SIZE, chunk_overlap=settings.CHUNK_OVERLAP)
    parsed_pages = ingester.parse_pdf(pdf_path)
    
    assert len(parsed_pages) == 2
    assert parsed_pages[0]["metadata"]["page"] == 1
    assert parsed_pages[1]["metadata"]["page"] == 2
    
    chunks = ingester.chunk_documents(parsed_pages)
    assert len(chunks) >= 2
    assert "Page 1" in chunks[0]["text"]

from src.ingestion.parser import DocumentIngester, UnsupportedFileTypeError

def test_parse_docx(tmp_path: Path):
    """Verify DOCX parsing extracts text elements with correct source metadata."""
    from docx import Document as DocxDocument

    docx_path = tmp_path / "sample.docx"
    doc = DocxDocument()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Hybrid retrieval combines dense and sparse search methods.")
    doc.save(str(docx_path))

    ingester = DocumentIngester()
    parsed = ingester.parse(docx_path)

    assert len(parsed) >= 2
    assert all(p["metadata"]["source"] == "sample.docx" for p in parsed)
    assert any("Hybrid retrieval" in p["text"] for p in parsed)


def test_parse_markdown(tmp_path: Path):
    """Verify Markdown parsing splits headers and body text into elements."""
    md_path = tmp_path / "notes.md"
    md_path.write_text("# Retrieval Notes\n\nBM25 is a classic sparse retrieval algorithm.")

    ingester = DocumentIngester()
    parsed = ingester.parse(md_path)

    assert len(parsed) >= 1
    assert any("BM25" in p["text"] for p in parsed)
    assert all(p["metadata"]["source"] == "notes.md" for p in parsed)


def test_parse_html(tmp_path: Path):
    """Verify HTML parsing strips tags and extracts visible text content."""
    html_path = tmp_path / "page.html"
    html_path.write_text(
        "<html><body><h1>RAG Systems</h1><p>Reranking improves precision.</p></body></html>"
    )

    ingester = DocumentIngester()
    parsed = ingester.parse(html_path)

    assert any("Reranking improves precision" in p["text"] for p in parsed)


def test_parse_plain_text_file(tmp_path: Path):
    """Verify .txt files route through the plain text parser."""
    txt_path = tmp_path / "readme.txt"
    txt_path.write_text("Chunking strategy affects retrieval quality significantly.")

    ingester = DocumentIngester()
    parsed = ingester.parse(txt_path)

    assert any("Chunking strategy" in p["text"] for p in parsed)


def test_parse_rejects_unsupported_extension(tmp_path: Path):
    """Verify unsupported file types raise a clear, catchable error instead of a cryptic crash."""
    bad_path = tmp_path / "malware.exe"
    bad_path.write_bytes(b"not a real document")

    ingester = DocumentIngester()
    with pytest.raises(UnsupportedFileTypeError):
        ingester.parse(bad_path)